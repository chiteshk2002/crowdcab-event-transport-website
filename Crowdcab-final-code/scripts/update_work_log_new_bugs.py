from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt, RGBColor


DOCX_PATH = Path(
    r"C:\Users\yashp\Desktop\MIT(Australia)\Semester 4\IT industry project\Project\crowdcab_work_log.docx"
)


def set_run(run, bold=False, size=10.5, color=None):
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run(run, bold=True, size=15, color=(0, 0, 0))
    return para


def add_body(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run(run, size=10.5, color=(35, 35, 35))
    return para


def add_two_col_table(doc, title, rows):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    table.rows[0].cells[0].text = title
    table.rows[0].cells[1].text = "Status / notes"
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                set_run(r, bold=True, size=10, color=(0, 0, 0))

    for bug_id, notes in rows:
        row = table.add_row().cells
        row[0].text = bug_id
        row[1].text = notes
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_run(r, size=9.5, color=(35, 35, 35))
    doc.add_paragraph()
    return table


def already_updated(doc):
    needle = "New testing bug report fixes"
    return any(needle in p.text for p in doc.paragraphs)


def main():
    doc = Document(DOCX_PATH)
    if already_updated(doc):
        print("Work log already contains the new testing bug report update.")
        return

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_heading(doc, "Update — 21 May 2026: New testing bug report fixes")
    add_body(
        doc,
        "I reviewed the new CrowdCab testing bug report and applied the valid fixes that fit the current prototype scope. "
        "The update focused on cleaning internal access, improving the Live Map pickup panel, removing misleading dashboard content, "
        "and preventing poor-quality pickup labels from appearing in recommendations.",
    )

    bug_rows = [
        (
            "NB-01",
            "Solved — admin and developer accounts are no longer presented as confusing public demo choices; internal pages now label role purpose.",
        ),
        (
            "NB-02",
            "Solved — the 1 Plan, 2 Choose, and 3 Ride controls were changed into non-clickable step indicators.",
        ),
        (
            "NB-03",
            "Solved — selected pickup cards on Live Map were made more compact so alternate pickup options stay visible.",
        ),
        (
            "NB-04",
            "Solved — internal demo access details were removed from the public login page.",
        ),
        (
            "NB-05",
            "Solved — internal staff accounts now require an INTERNAL_ACCESS_PASSWORD from the environment instead of email-only access.",
        ),
        (
            "NB-06",
            "Solved — admin/developer sessions no longer fall back to customer local trip history on My Trips.",
        ),
        (
            "NB-07",
            "Solved — dashboard wording now uses CrowdCab Web instead of the unclear app label.",
        ),
        (
            "NB-08",
            "Solved — the Payment Mix dashboard card was removed because payment is outside the current product scope.",
        ),
        (
            "NB-09",
            "Solved — pickup demand summaries now use confirmed pickup labels, so dashboard demand reflects changed pickup choices.",
        ),
        (
            "NB-10",
            "Solved — recommendation results now filter out raw Unnamed Service Road style labels when better named candidates exist.",
        ),
    ]
    add_two_col_table(doc, "Bug ID", bug_rows)

    add_heading(doc, "Files changed in this pass")
    files = [
        ".env.example",
        "app.py",
        "recommendation_engine.py",
        "templates/login.html",
        "templates/map.html",
        "templates/internal.html",
        "static/js/map.js",
        "static/js/my_trips.js",
        "static/js/dashboard.js",
        "static/css/styles.css",
    ]
    add_body(doc, ", ".join(files) + ".")

    add_heading(doc, "Validation completed")
    checks = [
        "Python compile check passed for app.py, recommendation_engine.py, and realtime_traffic.py.",
        "JavaScript syntax checks passed for all static/js files.",
        "Internal login without staff password is rejected.",
        "Booking dashboard no longer returns payment donut data.",
        "Brisbane Aquatic Centre recommendations no longer show Unnamed Service Road labels in the top results.",
    ]
    for item in checks:
        para = doc.add_paragraph()
        run = para.add_run(f"Check: {item}")
        set_run(run, size=10.5, color=(35, 35, 35))

    add_body(
        doc,
        "Remaining note: external open-data traffic checks can still fall back when network access is unavailable, but the app continues to operate safely with fallback scoring.",
    )

    doc.save(DOCX_PATH)
    print(f"Updated {DOCX_PATH}")


if __name__ == "__main__":
    main()
